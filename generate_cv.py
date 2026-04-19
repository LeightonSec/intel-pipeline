from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# Kill inherited spacing from the Normal style
normal = doc.styles['Normal']
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after  = Pt(0)

DARK       = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT     = RGBColor(0x0F, 0x3A, 0x6E)
LIGHT      = RGBColor(0x55, 0x55, 0x88)
BODY       = RGBColor(0x22, 0x22, 0x33)
SIDEBAR_BG = "F0F2F7"

# ── helpers ───────────────────────────────────────────────────────────────────

def _font(run, size, bold=False, italic=False, color=BODY):
    run.font.name    = "Calibri"
    run.font.size    = Pt(size)
    run.font.bold    = bold
    run.font.italic  = italic
    run.font.color.rgb = color

def add_rule(container, color="0F3A6E", width_pt=8):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(width_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def section_heading(container, text, size=9.5):
    p = container.add_paragraph()
    p.paragraph_format.space_before      = Pt(8)
    p.paragraph_format.space_after       = Pt(1)
    p.paragraph_format.page_break_before = False
    run = p.add_run(text.upper())
    _font(run, size, bold=True, color=ACCENT)
    add_rule(container, "0F3A6E", 6)

def body_para(container, text, size=10, space_after=4, indent=0):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _font(run, size)
    return p

def bullet_para(container, text, size=10, space_after=3):
    p = container.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    _font(run, size)
    return p

def job_header(container, title, company, dates):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(title)
    _font(r1, 10.5, bold=True, color=DARK)
    r2 = p.add_run(f"    {company}")
    _font(r2, 9.5, color=LIGHT)
    pd = container.add_paragraph()
    pd.paragraph_format.space_before = Pt(0)
    pd.paragraph_format.space_after  = Pt(3)
    _font(pd.add_run(dates), 9, italic=True, color=LIGHT)

def set_cell_shading(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_width(cell, width_cm):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def set_cell_padding(cell, top=0, right=0, bottom=0, left=0):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def remove_table_borders(table):
    # Table-level border removal
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"),  "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBdr.append(b)
    tblPr.append(tblBdr)
    # Cell-level border removal
    for row in table.rows:
        for cell in row.cells:
            tc    = cell._tc
            tcPr  = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBdr.append(b)
            tcPr.append(tcBdr)

def allow_row_break(row):
    # Explicitly remove any cantSplit element — Word default (absent = can split)
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for el in trPr.findall(qn("w:cantSplit")):
        trPr.remove(el)

def zero_para(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ── MASTER TABLE (2 rows × 2 cols) ───────────────────────────────────────────
# Row 0: header spanning full width
# Row 1: sidebar (left) | main narrative (right)

master = doc.add_table(rows=2, cols=2)
remove_table_borders(master)
allow_row_break(master.rows[0])
allow_row_break(master.rows[1])

# Merge row 0 into a single full-width header cell
HDR = master.cell(0, 0).merge(master.cell(0, 1))
set_cell_padding(HDR, top=0, right=0, bottom=80, left=0)

# Explicitly remove tblHeader flag — prevents Word treating row 0 as a
# repeating header that forces content rows onto page 2
tr0   = master.rows[0]._tr
trPr0 = tr0.get_or_add_trPr()
for el in trPr0.findall(qn("w:tblHeader")):
    trPr0.remove(el)
no_hdr = OxmlElement("w:tblHeader")
no_hdr.set(qn("w:val"), "0")
trPr0.append(no_hdr)

SB  = master.cell(1, 0)   # sidebar
MN  = master.cell(1, 1)   # main narrative

# Column widths (total usable = ~17cm)
set_cell_width(SB, 5.8)
set_cell_width(MN, 11.2)
set_cell_padding(SB, top=80, right=180, bottom=80, left=0)
set_cell_padding(MN, top=80, right=0,   bottom=80, left=140)
set_cell_shading(SB, SIDEBAR_BG)

# Zero out default empty paragraphs and suppress any page-break-before
for cell in [HDR, SB, MN]:
    for p in cell.paragraphs:
        zero_para(p)
        p.paragraph_format.page_break_before = False

# ── ROW 0 — HEADER ────────────────────────────────────────────────────────────

p_name = HDR.add_paragraph()
zero_para(p_name)
p_name.paragraph_format.space_after = Pt(2)
_font(p_name.add_run("LEIGHTON WILSON"), 24, bold=True, color=DARK)

p_tag = HDR.add_paragraph()
zero_para(p_tag)
p_tag.paragraph_format.space_after = Pt(3)
_font(p_tag.add_run(
    "Cybersecurity Professional  ·  SOC & Threat Detection  ·  Endpoint Security"
), 11, color=LIGHT)

p_contact = HDR.add_paragraph()
zero_para(p_contact)
p_contact.paragraph_format.space_after = Pt(4)
_font(p_contact.add_run(
    "+46 73 241 78 10  ·  leightonwsec@proton.me  ·  linkedin.com/in/leighton-w-7a220518b"
    "  ·  github.com/LeightonSec  ·  bastionprotocol.org  ·  Östersund, Sweden"
), 9, color=LIGHT)

add_rule(HDR, "0F3A6E", 14)

# ── ROW 1 LEFT — SIDEBAR ─────────────────────────────────────────────────────

section_heading(SB, "Certifications")
for icon, name, note in [
    ("✅", "CompTIA Security+",        "Awarded March 2026"),
    ("✅", "IBM Cybersecurity Analyst", "Professional Certificate"),
    ("✅", "ISC2 CC",                   "Entry-Level Certification"),
    ("✅", "Microsoft Learn",           "Azure & Security Fundamentals"),
    ("✅", "Linux Foundation Training", ""),
    ("🔄", "CompTIA Network+",          "In Progress — May 2026"),
    ("🔄", "Hack The Box Academy",      "Active Labs"),
]:
    p = SB.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{icon}  {name}")
    _font(r1, 9, bold=True)
    if note:
        r2 = p.add_run(f"\n    {note}")
        _font(r2, 8.5, color=LIGHT)

section_heading(SB, "Core Skills")
for skill in [
    "Endpoint Security & IAM",
    "Threat Detection & Analysis",
    "SIEM & Log Analysis",
    "Active Directory / Entra ID",
    "Network Security",
    "TCP/IP · DNS · VPN · Firewalls",
    "Linux & VPS Administration",
    "Python · Flask · PyShark",
    "LangGraph · Claude API",
    "Docker & GitHub",
    "ITIL Incident Management",
    "NHS / HIPAA Compliance",
    "Post-Quantum Cryptography",
]:
    p = SB.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.2)
    _font(p.add_run(f"▪  {skill}"), 9)

section_heading(SB, "Education")
p = SB.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(1)
_font(p.add_run("Electrical Installation\nApprenticeship"), 9, bold=True, color=DARK)
p2 = SB.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(1)
_font(p2.add_run("4-Year · City & Guilds"), 8.5, color=LIGHT)
p3 = SB.add_paragraph()
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after  = Pt(2)
_font(p3.add_run("Northern Regional College\n(formerly Larne Skills\nDevelopment Centre) · NI"), 8.5, color=LIGHT)

section_heading(SB, "Languages")
for lang, level in [
    ("English", "Native"),
    ("Swedish", "Beginner\nSFI enrolled May 2026\nDaily use at Systembolaget"),
]:
    p = SB.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    _font(p.add_run(f"{lang}\n"), 9, bold=True)
    _font(p.add_run(level), 8.5, color=LIGHT)

# ── ROW 1 RIGHT — MAIN NARRATIVE ─────────────────────────────────────────────

section_heading(MN, "Professional Summary")
body_para(MN,
    "Cybersecurity-focused IT Deployment Engineer with proven experience securing and managing "
    "critical infrastructure at scale across regulated healthcare environments. Holds CompTIA Security+ "
    "and is actively pursuing CompTIA Network+ (May 2026). Builds production-grade security tooling in "
    "Python — including a complete six-layer SOC toolkit with multi-agent AI orchestration, network "
    "packet analysis, and automated threat intelligence pipelines. Seeking a Junior SOC Analyst or IT "
    "Security role in Sweden or remote where practical endpoint security experience, software development "
    "capability, and genuine curiosity can contribute from day one.",
    size=9.5, space_after=2)

section_heading(MN, "Professional Experience")

job_header(MN,
    "IT Deployment Engineer",
    "Capita — South Eastern Health & Social Care Trust · NI",
    "Sep 2023 – Aug 2025")
for b in [
    "Delivered the largest healthcare IT modernisation in Northern Ireland's history — refreshing the full "
    "endpoint estate across 5 hospitals, all community health, mental health, admin and GP surgery sites, "
    "managing 10,000+ devices from procurement to secure disposal using SCCM and Dell Wyse Management Suite.",
    "Embedded directly with the cybersecurity team during a live NIST audit — personally remediating "
    "Sophos-flagged non-compliant endpoints, restoring compliant baselines at scale to maintain audit "
    "posture across a critical NHS environment.",
    "Held domain admin privileges across the full estate — enrolling and deploying endpoints through SCCM "
    "under ITIL-governed change management, ensuring every device met NHS security baselines before "
    "entering a clinical environment.",
    "Led go-live of Encompass EPR — the first deployment of its kind in Northern Ireland — embedded on the "
    "A&E floor and in critical care departments during cutover, providing sustained technical support for "
    "months post-launch with zero disruption to patient care.",
    "Operated as cross-functional technical liaison across cybersecurity, networking, IT support, data "
    "centre, and HQ — trusted with unsupervised access across all Trust sites, including launching a new "
    "radiology department with specialist clinical imaging equipment.",
]:
    bullet_para(MN, b, size=9.5)

job_header(MN,
    "Linux Infrastructure Operator",
    "Gala Games — Independent · Remote",
    "Nov 2020 – Present  —  concurrent with primary employment")
for b in [
    "Self-taught Linux administration from first principles — provisioning, hardening, and maintaining "
    "Ubuntu 20.04 VPS environments across DigitalOcean, Vultr, and RackNerd within a decentralised "
    "network of 50,000+ nodes.",
    "Managed containerised deployments using Docker with version control via GitHub — building practical "
    "understanding of distributed systems, blockchain infrastructure, and tokenisation mechanics autonomously.",
]:
    bullet_para(MN, b, size=9.5)

job_header(MN,
    "Market Research & Data Analyst",
    "Independent — Freelance · Remote",
    "Apr 2020 – Jul 2025  —  concurrent with primary employment")
bullet_para(MN,
    "Applied structured risk frameworks to time-critical decisions under pressure — developing "
    "pattern-recognition, evidence-based analysis, and rigorous documentation habits directly "
    "transferable to SOC work.",
    size=9.5)

section_heading(MN, "Portfolio — LeightonSec SOC Toolkit")
p_gh = MN.add_paragraph()
p_gh.paragraph_format.space_before = Pt(1)
p_gh.paragraph_format.space_after  = Pt(3)
_font(p_gh.add_run(
    "github.com/LeightonSec  —  Six-layer Security Operations toolkit, all projects live on GitHub."
), 9, italic=True, color=LIGHT)

for title, desc in [
    ("AI Firewall",       "LLM jailbreak detection with keyword pattern matching and Claude API semantic analysis. Risk scoring LOW/MEDIUM/HIGH with prompt injection hardening."),
    ("PCAP Analyser",     "Network packet capture analysis detecting DoS/DDoS, SYN floods, port scanning, ARP spoofing and C2 beaconing. AbuseIPDB threat intel integration."),
    ("Intel Pipeline",    "Automated threat intel ingestion from 12 RSS sources. Multi-agent LangGraph orchestration with Claude API summarisation, domain whitelist security layer, and twice-daily structured reporting."),
    ("Incident Tracker",  "SOC ticketing system with full REST API, ticket lifecycle management, escalation workflow and comment system."),
    ("Unified Dashboard", "Single pane of glass aggregating live data across all toolkit tools."),
    ("Intel Dashboard",   "Browser-based report viewer for intel pipeline output with search, filtering and source status."),
]:
    p = MN.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    _font(p.add_run(f"{title} — "), 9.5, bold=True, color=ACCENT)
    _font(p.add_run(desc), 9.5)

section_heading(MN, "Notable Project — Bastion Protocol")
p_bp = MN.add_paragraph()
p_bp.paragraph_format.space_before = Pt(2)
p_bp.paragraph_format.space_after  = Pt(2)
_font(p_bp.add_run("bastionprotocol.org  —  "), 9.5, italic=True, color=LIGHT)
_font(p_bp.add_run(
    "Founder and architect of Bastion — a sovereign, quantum-resistant communication protocol built on "
    "Bitcoin SV. Five-layer architecture implementing hybrid ECDH + ML-KEM (Kyber) post-quantum key "
    "encapsulation, on-chain guardian system, and multi-agent AI layer. One of the few hands-on "
    "implementations of NIST's post-quantum cryptography standard outside academic research. "
    "Whitepaper v0.1 complete."
), 9.5)

# ── FOOTER ────────────────────────────────────────────────────────────────────

p_avail = doc.add_paragraph()
p_avail.paragraph_format.space_before = Pt(6)
p_avail.paragraph_format.space_after  = Pt(0)
p_avail.alignment = WD_ALIGN_PARAGRAPH.CENTER
_font(p_avail.add_run(
    "Available immediately for Junior SOC Analyst and IT Security roles across Sweden and remote."
), 9.5, bold=True, color=ACCENT)

out = "/Users/leighton/Documents/Projects/intel-pipeline/Leighton_Wilson_CV.docx"
doc.save(out)
print(f"Saved: {out}")
